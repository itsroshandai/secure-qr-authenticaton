#!/usr/bin/env python3
"""
Script to generate Word document report for Secure QR Authentication System
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re
from io import StringIO

def markdown_to_docx(markdown_content, output_file):
    """Convert markdown content to Word document"""

    # Create a new document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Secure QR Authentication System - Project Report"
    doc.core_properties.author = "AI Assistant"
    doc.core_properties.subject = "Technical Project Report"

    # Split content into lines
    lines = markdown_content.split('\n')
    current_paragraph = []
    in_code_block = False
    code_content = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Handle headers
        if line.startswith('# '):
            # Add any pending paragraph
            if current_paragraph:
                p = doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []

            # Add heading
            heading = doc.add_heading(line[2:], level=1)
            heading.style = 'Heading 1'

        elif line.startswith('## '):
            if current_paragraph:
                p = doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []

            heading = doc.add_heading(line[3:], level=2)
            heading.style = 'Heading 2'

        elif line.startswith('### '):
            if current_paragraph:
                p = doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []

            heading = doc.add_heading(line[4:], level=3)
            heading.style = 'Heading 3'

        # Handle code blocks
        elif line.startswith('```'):
            if current_paragraph:
                p = doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []

            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                in_code_block = False
                # Add code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                runner = p.add_run(code_text)
                runner.font.name = 'Courier New'
                runner.font.size = Pt(10)

        elif in_code_block:
            code_content.append(line)

        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            if current_paragraph:
                p = doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []

            # Add bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')

        elif line.strip() == '':
            # Empty line - add current paragraph if any
            if current_paragraph:
                p = doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
        else:
            # Regular paragraph content
            current_paragraph.append(line)

        i += 1

    # Add any remaining content
    if current_paragraph:
        p = doc.add_paragraph(' '.join(current_paragraph))

    # Save the document
    doc.save(output_file)
    print(f"Word document saved as: {output_file}")

def main():
    # Read the markdown report
    with open('Secure_QR_Auth_Report.md', 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Generate Word document
    markdown_to_docx(markdown_content, 'Secure_QR_Auth_Report.docx')

if __name__ == '__main__':
    main()